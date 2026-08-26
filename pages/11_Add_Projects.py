
import io
import re
import random
import zipfile
from pathlib import Path

import pandas as pd
import streamlit as st

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ============================================================
# PATHS
# ============================================================

THIS_FILE = Path(__file__).resolve()
PROJECT_ROOT = (
    THIS_FILE.parent.parent
    if THIS_FILE.parent.name.lower() == "pages"
    else THIS_FILE.parent
)
DATA_DIR = PROJECT_ROOT / "data"
TRAINERS_FILE = DATA_DIR / "Trainers.xlsx"
CT_PROJECTS_FILE = DATA_DIR / "CTProjects.xlsx"


# ============================================================
# BASIC HELPERS
# ============================================================

def clean_text(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value).strip()


def normalize_text(value):
    text = clean_text(value).lower()
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9+#.]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def normalize_emp_id(value):
    text = clean_text(value).upper().replace(" ", "")
    digits = "".join(ch for ch in text if ch.isdigit())
    return f"CT{int(digits):04d}" if digits else text


def split_manual_topics(values):
    result = []
    seen = set()
    for value in values or []:
        topic = clean_text(value)
        if topic and topic.lower() not in seen:
            result.append(topic)
            seen.add(topic.lower())
    return result


# ============================================================
# READ EMPLOYEE ID / NAME FROM WORD PROFILE
# ============================================================

def iter_doc_text(doc):
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = clean_text(cell.text)
                if text:
                    yield text


def extract_profile_identity(doc):
    texts = list(iter_doc_text(doc))
    combined = "\n".join(texts)

    emp_id = ""
    name = ""

    # Strong Employee ID pattern first.
    match = re.search(
        r"(?:employee\s*id|emp\s*id)\s*[:\-]?\s*(CT\s*\d{3,6})",
        combined,
        flags=re.IGNORECASE,
    )
    if match:
        emp_id = normalize_emp_id(match.group(1))

    # Fallback: first CTxxxx token anywhere in the profile.
    if not emp_id:
        match = re.search(r"\bCT\s*\d{3,6}\b", combined, flags=re.IGNORECASE)
        if match:
            emp_id = normalize_emp_id(match.group(0))

    match = re.search(
        r"(?:^|\n)\s*name\s*[:\-]\s*([^\n\r]+)",
        combined,
        flags=re.IGNORECASE,
    )
    if match:
        name = clean_text(match.group(1))

    return emp_id, name


# ============================================================
# LOAD EXCEL DATA
# ============================================================

def _find_column(df, names):
    normalized = {normalize_text(c): c for c in df.columns}
    for name in names:
        if normalize_text(name) in normalized:
            return normalized[normalize_text(name)]
    return None


@st.cache_data(show_spinner=False)
def load_trainers_data(file_mtime):
    if not TRAINERS_FILE.exists():
        raise FileNotFoundError(f"Not found: {TRAINERS_FILE}")

    df = pd.read_excel(TRAINERS_FILE)
    if df.empty:
        return pd.DataFrame(columns=["_emp_id", "_skills"])

    emp_col = _find_column(
        df,
        ["Emp ID", "Employee ID", "EmpID", "EmployeeID"],
    )
    if emp_col is None:
        raise ValueError("Employee ID column not found in Trainers.xlsx.")

    skills_col = _find_column(df, ["Skills"])
    if skills_col is None:
        if len(df.columns) >= 15:
            skills_col = df.columns[14]  # Excel Column O
        else:
            raise ValueError(
                "Skills column not found and Trainers.xlsx has no Column O."
            )

    out = pd.DataFrame()
    out["_emp_id"] = df[emp_col].apply(normalize_emp_id)
    out["_skills"] = df[skills_col].apply(clean_text)
    return out


def get_trainer_skills(emp_id):
    mtime = TRAINERS_FILE.stat().st_mtime if TRAINERS_FILE.exists() else 0
    df = load_trainers_data(mtime)
    emp_id = normalize_emp_id(emp_id)
    row = df[df["_emp_id"] == emp_id]
    return "" if row.empty else clean_text(row.iloc[0]["_skills"])


@st.cache_data(show_spinner=False)
def load_project_pool(file_mtime):
    if not CT_PROJECTS_FILE.exists():
        raise FileNotFoundError(f"Not found: {CT_PROJECTS_FILE}")

    try:
        df = pd.read_excel(CT_PROJECTS_FILE, sheet_name="Projects")
    except ValueError:
        df = pd.read_excel(CT_PROJECTS_FILE)

    required = [
        "College / Client",
        "Location",
        "Domain / Subject Area",
    ]
    missing = [column for column in required if column not in df.columns]
    if missing:
        raise ValueError(
            "CTProjects.xlsx is missing: " + ", ".join(missing)
        )

    df = df.copy()
    for column in required:
        df[column] = df[column].apply(clean_text)

    df = df[
        (df["College / Client"] != "")
        & (df["Location"] != "")
        & (df["Domain / Subject Area"] != "")
    ].copy()

    return df.drop_duplicates(subset=required).reset_index(drop=True)


# ============================================================
# TECHNOLOGY / TOPIC MATCHING
# ============================================================

DSA_TERMS = (
    "dsa",
    "data structure",
    "data structures",
    "algorithm",
    "algorithms",
    "competitive coding",
    "competitive programming",
    "problem solving",
)


def is_dsa(text):
    return any(term in text for term in DSA_TERMS)


def is_c_subject(text):
    if "c++" in text or "c plus plus" in text or "c#" in text:
        return False
    return bool(
        re.search(r"\bc programming\b", text)
        or re.search(r"\bc language\b", text)
        or re.search(r"\bcore c\b", text)
        or re.search(r"\bc with dsa\b", text)
        or re.search(r"\bdsa using c\b", text)
        or re.search(r"\bdata structures using c\b", text)
        or re.fullmatch(r"c", text)
    )


def topics_from_trainer_skills(skills):
    text = normalize_text(skills)
    topics = ["C Programming"]  # allowed for every trainer

    def add(topic):
        if topic not in topics:
            topics.append(topic)

    if re.search(r"\bjava\b", text):
        add("Java")
        add("DSA using Java")

    if re.search(r"\bpython\b", text):
        add("Python")
        add("DSA using Python")

    if any(
        token in text
        for token in (
            "web development",
            "html",
            "css",
            "javascript",
            "react",
            "node",
            "express",
            "mern",
            "full stack",
            "fullstack",
            "django",
            "flask",
        )
    ):
        add("Web Development")

    if "mongodb" in text or "mongo db" in text:
        add("MongoDB")

    if any(token in text for token in ("sql", "dbms", "mysql", "database")):
        add("Database / SQL / DBMS")

    if any(
        token in text
        for token in (
            "artificial intelligence",
            "machine learning",
            "deep learning",
            "genai",
            "generative ai",
            "agentic ai",
            "prompt engineering",
            "n8n",
            "llm",
        )
    ):
        add("AI / ML / GenAI")

    if any(token in text for token in ("data science", "data analytics", "data analysis")):
        add("Data Science / Analytics")

    if "cyber security" in text or "cybersecurity" in text:
        add("Cyber Security")

    if any(token in text for token in ("aws", "azure", "gcp", "cloud computing")):
        add("Cloud")

    if "devops" in text:
        add("DevOps")

    if "power bi" in text or "powerbi" in text:
        add("Power BI")

    if "tableau" in text:
        add("Tableau")

    if "salesforce" in text:
        add("Salesforce")

    return topics


def project_matches_topic(subject, topic):
    s = normalize_text(subject)
    t = normalize_text(topic)

    if not s or not t:
        return False

    # Strong direct match for custom/new technology names.
    if t in s:
        return True

    # Known topic rules.
    if t in ("c", "c programming", "c language"):
        return is_c_subject(s)

    if t == "java":
        return bool(re.search(r"\bjava\b", s))

    if t in ("dsa using java", "java dsa"):
        return bool(re.search(r"\bjava\b", s)) and is_dsa(s)

    if t == "python":
        return bool(re.search(r"\bpython\b", s))

    if t in ("dsa using python", "python dsa"):
        return bool(re.search(r"\bpython\b", s)) and is_dsa(s)

    if t in ("web development", "web"):
        return any(
            token in s
            for token in (
                "web development",
                "html",
                "css",
                "javascript",
                "react",
                "node",
                "express",
                "mern",
                "full stack",
                "fullstack",
                "django",
                "flask",
            )
        )

    if t in ("mongodb", "mongo db"):
        return "mongodb" in s or "mongo db" in s

    if t in ("database sql dbms", "database", "sql", "dbms"):
        return any(token in s for token in ("sql", "dbms", "mysql", "database"))

    if t in ("ai ml genai", "ai", "machine learning", "genai"):
        return any(
            token in s
            for token in (
                "artificial intelligence",
                "machine learning",
                "deep learning",
                "genai",
                "generative ai",
                "agentic ai",
                "prompt engineering",
                "llm",
                "ai ",
            )
        )

    if t in ("data science analytics", "data science", "data analytics"):
        return any(token in s for token in ("data science", "data analytics", "data analysis"))

    if t in ("cyber security", "cybersecurity"):
        return "cyber security" in s or "cybersecurity" in s

    if t == "cloud":
        return any(token in s for token in ("aws", "azure", "gcp", "cloud"))

    if t == "devops":
        return "devops" in s

    if t == "power bi":
        return "power bi" in s or "powerbi" in s

    if t == "tableau":
        return "tableau" in s

    if t == "salesforce":
        return "salesforce" in s

    # Generic custom-topic fallback.
    words = [w for w in t.split() if len(w) > 2]
    return bool(words) and all(word in s for word in words)


def rows_for_topic(project_pool, topic, used_indexes=None):
    used_indexes = used_indexes or set()
    matches = []
    for index, row in project_pool.iterrows():
        if index in used_indexes:
            continue
        if project_matches_topic(row["Domain / Subject Area"], topic):
            matches.append(index)
    return matches


def pick_real_college_location(project_pool, used_pairs=None):
    used_pairs = used_pairs or set()
    base = project_pool[
        ["College / Client", "Location"]
    ].drop_duplicates().to_dict("records")

    random.shuffle(base)
    for row in base:
        pair = (
            clean_text(row["College / Client"]).lower(),
            clean_text(row["Location"]).lower(),
        )
        if pair not in used_pairs:
            return row

    return random.choice(base) if base else None


def synthesize_project(project_pool, topic, used_pairs=None):
    base = pick_real_college_location(project_pool, used_pairs)
    if not base:
        return None

    return {
        "College / Client": clean_text(base["College / Client"]),
        "Location": clean_text(base["Location"]),
        "Domain / Subject Area": clean_text(topic),
    }


def select_projects(project_pool, count, mode, trainer_skills="", manual_topics=None):
    count = int(count)
    manual_topics = split_manual_topics(manual_topics)
    used_indexes = set()
    used_pairs = set()
    rows = []
    synthesized_topics = []

    if mode == "Random from all CTProjects":
        take = min(count, len(project_pool))
        indexes = random.sample(list(project_pool.index), take)
        for index in indexes:
            row = project_pool.loc[index]
            rows.append(
                {
                    "College / Client": clean_text(row["College / Client"]),
                    "Location": clean_text(row["Location"]),
                    "Domain / Subject Area": clean_text(row["Domain / Subject Area"]),
                }
            )

    else:
        topics = (
            topics_from_trainer_skills(trainer_skills)
            if mode == "Based on Trainer Skills"
            else manual_topics
        )

        if not topics:
            return (
                pd.DataFrame(
                    columns=[
                        "S.No",
                        "College / Client",
                        "Location",
                        "Domain / Subject Area",
                    ]
                ),
                [],
                [],
            )

        # Manual mode: cycle across all user-entered technologies/topics.
        # Skill mode: randomize eligible trainer topics.
        topic_cycle = list(topics)
        if mode == "Based on Trainer Skills":
            random.shuffle(topic_cycle)

        slot = 0
        safety = 0
        while len(rows) < count and safety < count * max(10, len(topic_cycle) * 4):
            topic = topic_cycle[slot % len(topic_cycle)]
            slot += 1
            safety += 1

            matches = rows_for_topic(project_pool, topic, used_indexes)

            if matches:
                index = random.choice(matches)
                source = project_pool.loc[index]
                rows.append(
                    {
                        "College / Client": clean_text(source["College / Client"]),
                        "Location": clean_text(source["Location"]),
                        "Domain / Subject Area": clean_text(source["Domain / Subject Area"]),
                    }
                )
                used_indexes.add(index)
                used_pairs.add(
                    (
                        clean_text(source["College / Client"]).lower(),
                        clean_text(source["Location"]).lower(),
                    )
                )
                continue

            # If this topic has no matching source rows (or all are used),
            # create a new project using a real college/location from CTProjects.
            created = synthesize_project(project_pool, topic, used_pairs)
            if created:
                rows.append(created)
                used_pairs.add(
                    (
                        created["College / Client"].lower(),
                        created["Location"].lower(),
                    )
                )
                if topic not in synthesized_topics:
                    synthesized_topics.append(topic)

        topics_used = topics

    result = pd.DataFrame(
        rows,
        columns=[
            "College / Client",
            "Location",
            "Domain / Subject Area",
        ],
    )

    if result.empty:
        return (
            pd.DataFrame(
                columns=[
                    "S.No",
                    "College / Client",
                    "Location",
                    "Domain / Subject Area",
                ]
            ),
            [],
            synthesized_topics,
        )

    result.insert(0, "S.No", range(1, len(result) + 1))

    return (
        result,
        topics_used if mode != "Random from all CTProjects" else [],
        synthesized_topics,
    )


# ============================================================
# WORD TABLE / INSERTION
# ============================================================

def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_section_paragraph(doc, needle):
    needle = normalize_text(needle)
    for paragraph in doc.paragraphs:
        if needle in normalize_text(paragraph.text):
            return paragraph
    return None


def insert_paragraph_after(doc, reference_paragraph, text):
    paragraph = doc.add_paragraph()
    reference_paragraph._p.addnext(paragraph._p)
    run = paragraph.add_run(text)
    return paragraph, run


def xml_text(element):
    return "".join(node.text or "" for node in element.xpath(".//w:t")).strip()


def last_table_immediately_after(heading_paragraph):
    current = heading_paragraph._p.getnext()
    last_table = None

    while current is not None:
        tag = current.tag.split("}")[-1]

        if tag == "tbl":
            last_table = current
        elif tag == "p":
            if xml_text(current):
                break
        else:
            break

        current = current.getnext()

    return last_table


def set_table_borders(table, color="A6A6A6", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)

        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=60, bottom=55, end=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_projects_table(table):
    # IMPORTANT: no named Word style is used here.
    # This avoids "no style with name 'Table Grid'".
    set_table_borders(table)
    table.autofit = True

    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E97132")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row in table.rows[1:]:
        for col_index, cell in enumerate(row.cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if col_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)


def add_projects_table(doc, projects):
    table = doc.add_table(rows=1, cols=4)

    headers = [
        "S.No",
        "College / Client",
        "Location",
        "Domain / Subject Area",
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for _, row in projects.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["S.No"])
        cells[1].text = clean_text(row["College / Client"])
        cells[2].text = clean_text(row["Location"])
        cells[3].text = clean_text(row["Domain / Subject Area"])

    style_projects_table(table)
    return table


def insert_projects_into_document(doc, projects):
    # Remove the placeholder message anywhere in normal paragraphs.
    for paragraph in list(doc.paragraphs):
        if "project allocation not yet done" in normalize_text(paragraph.text):
            remove_paragraph(paragraph)

    training_heading = find_section_paragraph(doc, "Training Projects")

    if training_heading is None:
        training_heading = doc.add_paragraph()
        run = training_heading.add_run("Training Projects:")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(233, 113, 50)

    completed_heading = find_section_paragraph(doc, "Completed Projects")

    if completed_heading is None:
        completed_heading, run = insert_paragraph_after(
            doc,
            training_heading,
            "Completed Projects:",
        )
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(233, 113, 50)
        insert_after = completed_heading._p
    else:
        existing_table = last_table_immediately_after(completed_heading)
        insert_after = existing_table if existing_table is not None else completed_heading._p

    table = add_projects_table(doc, projects)
    insert_after.addnext(table._tbl)

    return doc


# ============================================================
# PROCESS ONE FILE
# ============================================================

def process_word_file(uploaded_file, project_count, mode, project_pool, manual_topics=None):
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    emp_id, trainer_name = extract_profile_identity(doc)

    trainer_skills = ""
    if mode == "Based on Trainer Skills" and emp_id:
        trainer_skills = get_trainer_skills(emp_id)

    projects, topics_used, synthesized = select_projects(
        project_pool=project_pool,
        count=project_count,
        mode=mode,
        trainer_skills=trainer_skills,
        manual_topics=manual_topics,
    )

    if projects.empty:
        return {
            "success": False,
            "filename": uploaded_file.name,
            "emp_id": emp_id,
            "trainer_name": trainer_name,
            "skills": trainer_skills,
            "message": "No projects could be generated.",
        }

    insert_projects_into_document(doc, projects)

    stream = io.BytesIO()
    doc.save(stream)

    stem = Path(uploaded_file.name).stem

    return {
        "success": True,
        "filename": uploaded_file.name,
        "output_name": f"{stem}_with_projects.docx",
        "output_bytes": stream.getvalue(),
        "emp_id": emp_id,
        "trainer_name": trainer_name,
        "skills": trainer_skills,
        "topics_used": topics_used,
        "synthesized_topics": synthesized,
        "projects": projects,
        "mode": mode,
    }


# ============================================================
# STREAMLIT PAGE
# ============================================================

def render_page():
    st.set_page_config(
        page_title="Add Projects",
        page_icon="➕",
        layout="wide",
    )

    st.title("➕ Add Projects to Trainer Profiles")
    st.caption(
        "Upload trainer Word profiles and add completed projects from CTProjects.xlsx."
    )

    left, right = st.columns(2)

    with left:
        if TRAINERS_FILE.exists():
            st.success("✅ Trainers.xlsx found")
        else:
            st.warning("⚠️ data/Trainers.xlsx not found")

    with right:
        if CT_PROJECTS_FILE.exists():
            try:
                pool_count = len(
                    load_project_pool(CT_PROJECTS_FILE.stat().st_mtime)
                )
                st.success(f"✅ CTProjects.xlsx found — {pool_count} projects")
            except Exception:
                st.success("✅ CTProjects.xlsx found")
        else:
            st.error("❌ data/CTProjects.xlsx not found")

    st.markdown("### 1. Upload Trainer Word Files")

    uploaded_files = st.file_uploader(
        "Upload one or more .docx trainer profiles",
        type=["docx"],
        accept_multiple_files=True,
    )

    st.markdown("### 2. Project Selection")

    col1, col2 = st.columns(2)

    with col1:
        number_of_projects = st.number_input(
            "Number of projects to add to EACH trainer",
            min_value=1,
            max_value=30,
            value=5,
            step=1,
        )

    with col2:
        selection_mode = st.radio(
            "Technology Selection",
            options=[
                "Based on Trainer Skills",
                "Add Skills for Projects",
                "Random from all CTProjects",
            ],
            index=0,
        )

    manual_topics = []

    if selection_mode == "Based on Trainer Skills":
        st.info(
            "**Java → Java + DSA using Java**  |  "
            "**Python → Python + DSA using Python**  |  "
            "**C Programming is allowed for every trainer**. "
            "Skills are read from Trainers.xlsx Column O / Skills."
        )

    elif selection_mode == "Add Skills for Projects":
        st.markdown("### 3. Add Skills / Technologies")

        st.caption(
            "Enter one skill or technology in each box. "
            "Click **Add One More Skill / Technology** for another box."
        )

        if "ap_topic_count" not in st.session_state:
            st.session_state.ap_topic_count = 1

        for index in range(st.session_state.ap_topic_count):
            value = st.text_input(
                f"Skill / Technology {index + 1}",
                key=f"ap_topic_{index}",
                placeholder="Example: Python / Java / React Native / Snowflake",
            )
            if clean_text(value):
                manual_topics.append(clean_text(value))

        add_col, remove_col, reset_col = st.columns(3)

        with add_col:
            if st.button(
                "➕ Add One More Skill / Technology",
                use_container_width=True,
            ):
                st.session_state.ap_topic_count += 1
                st.rerun()

        with remove_col:
            if st.button(
                "➖ Remove Last",
                use_container_width=True,
                disabled=st.session_state.ap_topic_count <= 1,
            ):
                last_index = st.session_state.ap_topic_count - 1
                st.session_state.pop(f"ap_topic_{last_index}", None)
                st.session_state.ap_topic_count -= 1
                st.rerun()

        with reset_col:
            if st.button("↺ Reset Skills", use_container_width=True):
                for index in range(st.session_state.ap_topic_count):
                    st.session_state.pop(f"ap_topic_{index}", None)
                st.session_state.ap_topic_count = 1
                st.rerun()

        if manual_topics:
            st.success("Projects will use: " + ", ".join(manual_topics))
        else:
            st.warning("Enter at least one skill / technology.")

        if len(manual_topics) > int(number_of_projects):
            st.warning(
                "You entered more technologies than the number of projects. "
                "Increase the project count if you want at least one project for every technology."
            )

        st.caption(
            "If a technology is not present in CTProjects.xlsx, "
            "the app uses a real College / Client and Location from CTProjects.xlsx "
            "and uses your entered technology as the Domain / Subject Area."
        )

    else:
        st.info(
            "Random mode selects from the complete CTProjects.xlsx project pool."
        )

    if "add_projects_results" not in st.session_state:
        st.session_state.add_projects_results = []

    disabled = (
        not uploaded_files
        or not CT_PROJECTS_FILE.exists()
        or (
            selection_mode == "Based on Trainer Skills"
            and not TRAINERS_FILE.exists()
        )
        or (
            selection_mode == "Add Skills for Projects"
            and not manual_topics
        )
    )

    if st.button(
        "🚀 Generate and Add Projects",
        type="primary",
        use_container_width=True,
        disabled=disabled,
    ):
        try:
            project_pool = load_project_pool(CT_PROJECTS_FILE.stat().st_mtime)
            results = []
            progress = st.progress(0)

            for position, uploaded_file in enumerate(uploaded_files, start=1):
                result = process_word_file(
                    uploaded_file=uploaded_file,
                    project_count=number_of_projects,
                    mode=selection_mode,
                    project_pool=project_pool,
                    manual_topics=manual_topics,
                )
                results.append(result)
                progress.progress(int(position / len(uploaded_files) * 100))

            st.session_state.add_projects_results = results

            success_count = sum(1 for r in results if r.get("success"))
            st.success(
                f"Completed: {success_count} of {len(results)} file(s) processed."
            )

        except Exception as exc:
            st.exception(exc)

    results = st.session_state.add_projects_results

    if not results:
        return

    st.markdown("---")
    st.markdown("### 4. Processed Files")

    successful = [r for r in results if r.get("success")]

    if successful:
        zip_stream = io.BytesIO()

        with zipfile.ZipFile(
            zip_stream,
            "w",
            zipfile.ZIP_DEFLATED,
        ) as archive:
            for result in successful:
                archive.writestr(
                    result["output_name"],
                    result["output_bytes"],
                )

        st.download_button(
            "📦 Download All Updated Profiles as ZIP",
            data=zip_stream.getvalue(),
            file_name="Trainer_Profiles_With_Projects.zip",
            mime="application/zip",
            use_container_width=True,
        )

    for index, result in enumerate(results, start=1):
        title = (
            result.get("trainer_name")
            or result.get("filename")
            or f"File {index}"
        )

        with st.expander(
            f"{index}. {title}",
            expanded=len(results) <= 5,
        ):
            if not result.get("success"):
                st.error(result.get("message", "Processing failed."))
                continue

            st.write("**Employee ID:**", result.get("emp_id") or "Not detected")
            st.write("**Selection mode:**", result.get("mode"))

            if result.get("skills"):
                st.write("**Skills from Trainers.xlsx:**", result["skills"])

            if result.get("topics_used"):
                st.write(
                    "**Technologies used for selection:**",
                    ", ".join(result["topics_used"]),
                )

            if result.get("synthesized_topics"):
                st.warning(
                    "Not found / insufficient in CTProjects.xlsx; "
                    "generated using a real college/location from CTProjects.xlsx: "
                    + ", ".join(result["synthesized_topics"])
                )

            st.dataframe(
                result["projects"][
                    [
                        "S.No",
                        "College / Client",
                        "Location",
                        "Domain / Subject Area",
                    ]
                ],
                hide_index=True,
                use_container_width=True,
            )

            st.download_button(
                f"⬇️ Download {result['output_name']}",
                data=result["output_bytes"],
                file_name=result["output_name"],
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                key=f"download_{index}_{result['output_name']}",
                use_container_width=True,
            )


if __name__ == "__main__":
    render_page()
