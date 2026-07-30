import os
import io
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from modules.project_extractor import ProjectExtractor


class WordGenerator:

    def __init__(self, template_path, trainer_data):

        self.template_path = template_path

        # Normalize Excel column names
        self.trainer = {
            str(k).strip(): v
            for k, v in trainer_data.items()
        }


        # Add missing optional fields
        default_fields = [

            "Professional Highlights",

            "LeetCode Profile Link",

            "No of Problems Done",

            "Other Profiles",

            "Training Expertise",

            "Competitive Programming",

            "Technical Expertise",

            "Certifications"

        ]


        for field in default_fields:

            if field not in self.trainer:

                self.trainer[field] = ""


        project_file = (

            "data/Projects.xlsx"

            if os.path.exists("data/Projects.xlsx")

            else "Projects.xlsx"

        )


        self.project_extractor = ProjectExtractor(
            project_file
        )

    ###########################################################
    # Table Formatting
    ###########################################################

    def set_cell_background(self, cell, fill_color):

        tc_pr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)

        tc_pr.append(shd)



    def set_table_borders(self, table):

        tbl_pr = table._tbl.tblPr

        tbl_borders = OxmlElement("w:tblBorders")

        for border_name in [
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ]:

            border = OxmlElement(
                f"w:{border_name}"
            )

            border.set(
                qn("w:val"),
                "single"
            )

            border.set(
                qn("w:sz"),
                "4"
            )

            border.set(
                qn("w:space"),
                "0"
            )

            border.set(
                qn("w:color"),
                "D3D3D3"
            )

            tbl_borders.append(border)


        tbl_pr.append(tbl_borders)

    ###########################################################
    # Placeholder Replacement
    ###########################################################

    def replace_placeholders(self, doc):


        # Normal paragraphs

        for paragraph in doc.paragraphs:

            self._replace_in_paragraph(
                paragraph
            )


        # Tables

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        self._replace_in_paragraph(
                            paragraph
                        )


        self.replace_image_placeholders(
            doc
        )



    def _replace_in_paragraph(self, paragraph):

        """
        Replaces placeholders safely.
        Handles Word split runs.
        """


        text = paragraph.text


        if not text:

            return



        original_text = text



        for key, value in self.trainer.items():


            if key == "Image":

                continue



            placeholder = (
                "{{"
                + str(key)
                + "}}"
            )



            if placeholder in text:


                if value is None:

                    value = ""


                text = text.replace(

                    placeholder,

                    str(value)

                )



        # Update only if changed

        if text != original_text:

            paragraph.text = text



    ###########################################################
    # Image Handling
    ###########################################################

    def replace_image_placeholders(self, doc):

        image_path = self.trainer.get(
            "Image",
            ""
        )


        if not image_path or not os.path.exists(image_path):

            self._clear_image_placeholder_text(
                doc
            )

            return



        for paragraph in doc.paragraphs:

            self._insert_image(
                paragraph,
                image_path
            )



        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        self._insert_image(
                            paragraph,
                            image_path
                        )



    def _insert_image(self, paragraph, image_path):


        if "{{Image}}" not in paragraph.text:

            return



        paragraph.text = paragraph.text.replace(

            "{{Image}}",

            ""

        )



        try:

            paragraph.add_run().add_picture(

                image_path,

                width=Inches(1.6)

            )


        except Exception:

            pass



    def _clear_image_placeholder_text(self, doc):


        for paragraph in doc.paragraphs:


            if "{{Image}}" in paragraph.text:


                paragraph.text = paragraph.text.replace(

                    "{{Image}}",

                    ""

                )



        for table in doc.tables:


            for row in table.rows:


                for cell in row.cells:


                    for paragraph in cell.paragraphs:


                        if "{{Image}}" in paragraph.text:


                            paragraph.text = paragraph.text.replace(

                                "{{Image}}",

                                ""

                            )
    ###########################################################
    # Projects Table
    ###########################################################

    def append_projects_table(
        self,
        doc,
        projects_list,
        section_title
    ):

        p = doc.add_paragraph()

        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(section_title)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)

        table = doc.add_table(
            rows=1,
            cols=4
        )

        self.set_table_borders(table)

        hdr_cells = table.rows[0].cells
        headers = [
            "S.No",
            "College / Client",
            "Location",
            "Domain / Subject Area"
        ]

        widths = [
            Inches(0.6),
            Inches(2.5),
            Inches(1.2),
            Inches(2.5)
        ]

        for i, header in enumerate(headers):

            hdr_cells[i].text = header

            run = hdr_cells[i].paragraphs[0].runs[0]

            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(255, 255, 255)

            self.set_cell_background(
                hdr_cells[i],
                "1F4E78"
            )

            hdr_cells[i].width = widths[i]

        for item in projects_list:

            row = table.add_row().cells

            

            row[0].text = str(item.get("S.No", ""))
            row[1].text = str(item.get("College", ""))
            row[2].text = str(item.get("Place", ""))
            row[3].text = str(item.get("Subject", ""))
            

            for i in range(4):

                row[i].width = widths[i]

                if row[i].paragraphs[0].runs:

                    run = row[i].paragraphs[0].runs[0]

                    run.font.name = "Calibri"
                    run.font.size = Pt(9)

    ###########################################################
    # Build Document
    ###########################################################

    def _build_document(self):

        if not os.path.exists(self.template_path):

            raise FileNotFoundError(
                f"Template not found : {self.template_path}"
            )

        doc = Document(self.template_path)

        # Replace placeholders
        self.replace_placeholders(doc)

        

        #######################################################
        # Projects
        #######################################################

        emp_id = (
            self.trainer.get("Emp ID")
            or self.trainer.get("Employee ID")
            or ""
        )

        emp_id = str(emp_id).strip()


        if emp_id.isdigit():

            emp_id = f"CT{int(emp_id):04d}"

        if not emp_id:

            print(
                "[WARNING] Employee ID not found.",
                file=sys.stderr
            )

        project_data = self.project_extractor.get_trainer_projects(emp_id)

        if isinstance(project_data, dict):

            ongoing = project_data.get("ongoing", [])

            completed = project_data.get("completed", [])

        else:

            ongoing = []
            completed = []

        if not ongoing and not completed:

            p = doc.add_paragraph()

            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

            run = p.add_run(
                "Project allocation not yet done."
            )

            run.italic = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"

        else:

            if ongoing:

                self.append_projects_table(
                    doc,
                    ongoing,
                    "Ongoing Projects:"
                )

            if completed:


                # Fill minimum 5 completed projects
                if len(completed) < 5:

                    required = 5 - len(completed)


                    extra_projects = self.project_extractor.get_random_completed_projects(
                        emp_id,
                        required
                    )


                    for p in extra_projects:

                        completed.append(p)



                self.append_projects_table(
                    doc,
                    completed,
                    "Completed Projects:"
                )

        # Final cleanup
        

        return doc

    ###########################################################
    # Save Document
    ###########################################################

    def generate(self, output_path):

        doc = self._build_document()

        doc.save(output_path)

    ###########################################################
    # Generate Bytes
    ###########################################################

    def generate_bytes(self):

        doc = self._build_document()

        buffer = io.BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        return buffer.getvalue()